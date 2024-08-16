import speedtest as st
import time
from docx import Document

def Speed_Test():
    
    Fhora =  time.strftime("%d-%m-%Y %H:%M:%S", time.localtime())
    print("Fecha y Hora ", Fhora)
    
    test = st.Speedtest()
    down_speed = test.download()
    down_speed = round(down_speed / 10**6, 2)
    print("Dowload Speed in Mbps: ", down_speed)

    up_speed = test.upload()
    up_speed = round(up_speed / 10**6, 2)
    print("Upload Speed in Mbps: ", up_speed)

    ping = test.results.ping
    print("Ping: ", ping)

    document = Document()

    document.add_heading('Test de Velocidad de Internet',0)

    document.add_paragraph('Fecha y Hora: ', style='Subtitle')
    document.add_paragraph(str(Fhora))

    document.add_paragraph('La velocidad de Internet de subida es de: ', style='Subtitle')
    document.add_paragraph(str(down_speed))

    document.add_paragraph('La velocidad de Internet de bajada es de: ', style='Subtitle')
    document.add_paragraph(str(up_speed))



    document.save("TestdeVelocidad.docx")

Speed_Test()
